from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_heading(paragraph, text, level=1):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after  = Pt(4)

def add_section_heading(doc, label, title):
    p = doc.add_paragraph()
    set_heading(p, f"{label}  {title}", level=2)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_math(doc, formula):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(formula)
    r.font.size   = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Research & Documentation Notes")
tr.font.size  = Pt(22)
tr.bold       = True
tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run("Bivariate Analysis in Credit Risk Modelling")
sr.font.size  = Pt(16)
sr.bold       = True
sr.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("FNB DataQuest 2026  |  Credit Scorecard Project\n"
             "Senior Data Scientist — Credit Risk Analysis Division").font.size = Pt(10)

add_rule(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION A — What is it?
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "A.", "What Is It?")
add_rule(doc)

add_body(doc,
    "Bivariate analysis is the simultaneous examination of exactly two variables to understand "
    "the nature, direction, and strength of the relationship between them. Where univariate analysis "
    "describes a single variable's distribution in isolation, bivariate analysis reveals how two "
    "variables co-vary — whether they move together, oppose each other, or are independent.")

add_body(doc,
    "It exists because credit default is not caused by a single factor in isolation. A borrower "
    "who earns R150 000 per year is not inherently high or low risk; the risk depends on whether "
    "that income is paired with R50 000 of debt (low risk) or R400 000 of debt (high risk). "
    "Bivariate analysis quantifies these pairwise dependencies, providing the empirical foundation "
    "for feature selection and, crucially, for the Weight of Evidence (WoE) transformation that "
    "underpins traditional credit scorecard development.")

add_body(doc,
    "Problem it solves: Univariate analysis cannot tell us which variables are predictive of "
    "default — it only tells us what each variable looks like individually. Bivariate analysis "
    "bridges the gap between raw data exploration and model building by identifying which features "
    "carry genuine information about the target (default / no default), quantifying the strength "
    "of that relationship, and detecting multicollinearity between predictors that would destabilise "
    "a logistic regression model.")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION B — How does it work?
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "B.", "How Does It Work?")
add_rule(doc)

add_body(doc,
    "The method applied depends on the types of the two variables being compared. "
    "In credit risk, the most important bivariate relationship is always a feature against "
    "the binary target (Default = 1 / No Default = 0). We also examine feature-feature "
    "relationships to detect redundancy.")

p = doc.add_paragraph()
set_heading(p, "1. Continuous Feature vs. Binary Target", level=3)

add_body(doc, "Point-biserial correlation coefficient (rpb):")
add_math(doc, "rpb = (M₁ − M₀) / Sₙ  ×  √(n₁·n₀ / n²)")
add_body(doc,
    "where M₁ and M₀ are the means of the feature in the default and non-default groups, "
    "Sₙ is the pooled standard deviation, n₁ and n₀ are group sizes, and n = n₁ + n₀. "
    "rpb ranges from −1 to +1. A higher absolute value indicates stronger separation between "
    "defaulters and non-defaulters on that feature.")

add_body(doc,
    "Visualisation: side-by-side box plots or overlapping KDE plots for the feature, "
    "coloured by default status. A large separation between the two distributions indicates "
    "high discriminatory power.")

p = doc.add_paragraph()
set_heading(p, "2. Continuous Feature vs. Continuous Feature", level=3)

add_body(doc, "Pearson correlation coefficient (r) — for linear relationships:")
add_math(doc, "r = Σ(xᵢ − x̄)(yᵢ − ȳ) / √[Σ(xᵢ − x̄)² · Σ(yᵢ − ȳ)²]")

add_body(doc, "Spearman rank correlation coefficient (ρ) — for monotonic relationships:")
add_math(doc, "ρ = 1 − (6 Σdᵢ²) / (n(n²−1))")
add_body(doc,
    "where dᵢ is the difference in ranks of each paired observation. Spearman is preferred "
    "in credit data because it is robust to outliers and non-normality — both common in "
    "financial variables. |ρ| > 0.7 between two predictors signals problematic multicollinearity "
    "for logistic regression.")

p = doc.add_paragraph()
set_heading(p, "3. Categorical Feature vs. Binary Target", level=3)

add_body(doc, "Chi-square test of independence (χ²):")
add_math(doc, "χ² = Σ (Oᵢⱼ − Eᵢⱼ)² / Eᵢⱼ")
add_body(doc,
    "where Oᵢⱼ is the observed count and Eᵢⱼ = (row total × column total) / n is the expected "
    "count under independence. A significant χ² (p < 0.05) indicates the categorical variable "
    "is associated with default status. The effect size is measured by Cramér's V.")

add_body(doc, "Cramér's V (effect size for chi-square):")
add_math(doc, "V = √(χ² / (n · min(r−1, c−1)))")
add_body(doc,
    "V = 0: no association; V = 1: perfect association. V > 0.3 is considered a moderate "
    "to strong association in credit risk applications.")

p = doc.add_paragraph()
set_heading(p, "4. Information Value (IV) and Weight of Evidence (WoE) — Credit-Specific Method", level=3)

add_body(doc,
    "IV and WoE are the industry-standard bivariate tools in credit scorecard development "
    "(originated by Fair Isaac / FICO). They measure how strongly a binned feature separates "
    "defaulters from non-defaulters.")

add_body(doc, "Weight of Evidence for bin i:")
add_math(doc, "WoEᵢ = ln(Distribution of Events_i / Distribution of Non-Events_i)")
add_math(doc, "     = ln((n_events_i / N_events) / (n_non_events_i / N_non_events))")

add_body(doc, "Information Value for a feature:")
add_math(doc, "IV = Σᵢ (Distribution of Events_i − Distribution of Non-Events_i) × WoEᵢ")

add_body(doc,
    "IV interpretation thresholds (Siddiqi, 2006):\n"
    "  IV < 0.02  →  Useless predictor (exclude)\n"
    "  0.02 ≤ IV < 0.10  →  Weak predictor\n"
    "  0.10 ≤ IV < 0.30  →  Medium predictor\n"
    "  0.30 ≤ IV < 0.50  →  Strong predictor\n"
    "  IV ≥ 0.50  →  Suspicious (possible data leakage — investigate)")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION C — Example in Credit Risk
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "C.", "Example in Credit Risk")
add_rule(doc)

add_body(doc,
    "Below are three worked examples of bivariate analysis from a retail lending dataset, "
    "covering each relationship type discussed in Section B.")

p = doc.add_paragraph()
set_heading(p, "Example 1 — Credit Bureau Score vs. Default (Continuous vs. Binary, WoE/IV)", level=3)
add_body(doc,
    "The credit bureau score is binned into five bands. The WoE and event rates for each bin are:")

woe_rows = [
    ("300 – 499 (Very Poor)",  "18.4%", "+0.82", "Predict Default — high WoE"),
    ("500 – 579 (Poor)",       "12.1%", "+0.41", "Above-average risk"),
    ("580 – 669 (Fair)",       "7.3%",  "+0.05", "Near population average"),
    ("670 – 739 (Good)",       "3.8%",  "−0.38", "Below-average risk"),
    ("740 – 850 (Excellent)",  "1.2%",  "−0.91", "Predict No Default — low WoE"),
]
for band, dr, woe, interp in woe_rows:
    add_bullet(doc, f"Default rate {dr} | WoE = {woe} | {interp}", bold_prefix=f"{band}: ")

add_body(doc,
    "The IV for credit bureau score = 0.38, classifying it as a strong predictor. "
    "The monotonically decreasing WoE (from +0.82 to −0.91) confirms the expected direction: "
    "higher scores → lower default probability. This monotonicity is a key quality check — "
    "non-monotonic WoE patterns indicate the binning needs refinement.")

p = doc.add_paragraph()
set_heading(p, "Example 2 — Employment Type vs. Default (Categorical vs. Binary, Chi-Square)", level=3)
add_body(doc,
    "A chi-square test is applied to employment type (Salaried, Self-Employed, Unemployed, Retired) "
    "against default status. χ²(3) = 142.6, p < 0.001, Cramér's V = 0.31. The association is "
    "statistically significant and moderately strong. Cross-tabulation shows self-employed borrowers "
    "default at 14.2% versus 6.1% for salaried borrowers — the largest differential. Decision: "
    "include employment type with WoE encoding after grouping 'Retired' and 'Other' due to small cell counts.")

p = doc.add_paragraph()
set_heading(p, "Example 3 — Income vs. DTI Ratio (Continuous vs. Continuous, Multicollinearity Check)", level=3)
add_body(doc,
    "Spearman correlation between log(annual income) and DTI ratio is ρ = −0.61. This is a moderate "
    "negative correlation (higher income tends to reduce the debt-to-income burden), but not strong "
    "enough (|ρ| < 0.70) to warrant dropping either variable. Both are retained. However, this "
    "correlation is flagged in the Variance Inflation Factor (VIF) check during logistic regression "
    "fitting. If VIF exceeds 5 for either variable, one would be dropped or an interaction term considered.")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION D — Strengths and Weaknesses
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "D.", "Strengths and Weaknesses")
add_rule(doc)

p = doc.add_paragraph()
set_heading(p, "Strengths", level=3)

strengths = [
    ("Identifies predictive features before modelling. ",
     "IV analysis provides a clear, rank-ordered list of how informative each feature is with "
     "respect to the target. This guides feature selection without fitting a model — reducing the "
     "risk of overfitting that comes from including too many weak predictors."),
    ("WoE enables direct input to logistic regression. ",
     "WoE-transformed features make logistic regression coefficients approximately equal across "
     "features (all on the same log-odds scale), improving numerical stability and interpretability. "
     "The linear relationship between WoE and log-odds is mathematically exact when the logistic "
     "regression model is correctly specified."),
    ("Handles non-linearity through binning. ",
     "By analysing the relationship within bins rather than assuming a global linear form, "
     "WoE/IV captures non-linear and non-monotonic relationships that would be invisible to "
     "Pearson correlation."),
    ("Detects multicollinearity early. ",
     "Spearman correlation matrices between all candidate features, computed before modelling, "
     "flag redundant variable pairs that would destabilise logistic regression coefficients."),
    ("Regulatory and industry standard. ",
     "WoE/IV analysis is explicitly referenced in Basel II/III model development guidelines and "
     "is standard practice in South African retail banking scorecard development."),
    ("Interpretable output. ",
     "WoE tables and IV rankings can be presented directly to credit policy teams and model "
     "validators — they do not require statistical expertise to understand."),
]
for bold, rest in strengths:
    add_bullet(doc, rest, bold_prefix=bold)

p = doc.add_paragraph()
set_heading(p, "Weaknesses", level=3)

weaknesses = [
    ("Only pairwise — ignores three-way and higher-order interactions. ",
     "A feature with low IV in isolation may become highly informative in combination with another "
     "feature. Bivariate analysis cannot detect this. Only model-based methods (classification trees, "
     "logistic regression with interaction terms) reveal these patterns."),
    ("Correlation ≠ causation. ",
     "A strong WoE signal for postcode might reflect socioeconomic confounding rather than a causal "
     "relationship with default. Bivariate analysis has no mechanism to distinguish correlation from "
     "causation — subject-matter judgement is always required."),
    ("IV is sensitive to bin granularity. ",
     "Too many fine-grained bins inflate IV by overfitting the training data. Bins with fewer than "
     "5% of observations produce unreliable WoE estimates. Monotonicity constraints and minimum "
     "bin-size rules are essential safeguards."),
    ("Pearson correlation misses non-linear relationships. ",
     "Two variables can have Pearson r ≈ 0 and still be strongly dependent (e.g., a U-shaped "
     "relationship). Spearman and WoE analysis partially address this, but visual inspection of "
     "scatter plots remains necessary."),
    ("Computationally intensive for high-dimensional data. ",
     "Pairwise bivariate analysis scales as O(p²) where p is the number of features. For datasets "
     "with hundreds of features, full pairwise analysis is expensive and produces results that are "
     "difficult to interpret without automated filtering."),
]
for bold, rest in weaknesses:
    add_bullet(doc, rest, bold_prefix=bold)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION E — Comparison to Logistic Regression
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "E.", "Comparison to Logistic Regression")
add_rule(doc)

add_body(doc,
    "Bivariate analysis and logistic regression occupy different stages of the modelling pipeline. "
    "Bivariate analysis is a precursor and enabler of logistic regression — not a competing method. "
    "The key distinctions are as follows.")

comparisons = [
    ("Scope of relationships modelled",
     "Logistic Regression: Models all features jointly, estimating coefficients that represent the "
     "partial effect of each feature after controlling for all others. Can detect suppressor variables "
     "(a feature with low marginal IV that becomes important after controlling for correlated features).",
     "Bivariate Analysis: One-at-a-time feature-target relationships. Cannot control for confounders. "
     "A feature that appears informative in bivariate analysis may be irrelevant once a correlated "
     "feature is included in the model."),
    ("Output",
     "Logistic Regression: A predictive score P(Default = 1 | X₁, …, Xₚ); odds ratios; AUC; Gini "
     "coefficient; calibration statistics. Directly usable for application scoring.",
     "Bivariate Analysis: IV rankings, WoE tables, correlation matrices, chi-square p-values. "
     "Not predictive on their own — serve as inputs to feature selection and transformation decisions."),
    ("Non-linearity",
     "Logistic Regression: Assumes a linear relationship between each WoE-transformed feature and "
     "log-odds. Non-linearity must be pre-encoded through binning/WoE.",
     "Bivariate Analysis: WoE/IV captures non-linear feature-target relationships within bins "
     "without any model assumptions. This is precisely why WoE transformation is applied before "
     "logistic regression."),
    ("Stability and regularisation",
     "Logistic Regression: Coefficient estimates are influenced by all correlated features "
     "simultaneously. High multicollinearity inflates standard errors and destabilises coefficients. "
     "L2 regularisation (ridge) mitigates this.",
     "Bivariate Analysis: Pearson/Spearman correlation between features directly identifies "
     "multicollinear pairs before they enter the model, allowing pre-emptive feature removal."),
    ("Regulatory standing",
     "Logistic Regression: The primary model submitted for regulatory approval (PD model under Basel). "
     "Coefficients and scorecard weights are the regulated artefact.",
     "Bivariate Analysis: Supporting documentation for model development. Regulators expect IV "
     "tables and WoE plots as evidence that features were understood and selected rigorously."),
    ("When preferred",
     "Logistic Regression: Always preferred as the final scoring model. No bivariate method produces "
     "a deployable credit score.",
     "Bivariate Analysis: Always performed before logistic regression as a non-negotiable EDA "
     "and feature engineering step. The WoE output directly feeds into the logistic regression "
     "as transformed inputs."),
]

for topic, lr_text, ba_text in comparisons:
    p = doc.add_paragraph()
    set_heading(p, topic, level=3)
    add_bullet(doc, lr_text, bold_prefix="Logistic Regression:  ")
    add_bullet(doc, ba_text, bold_prefix="Bivariate Analysis:   ")

add_body(doc,
    "In the credit scorecard development process, bivariate analysis and logistic regression are "
    "not alternatives — they are sequential steps in the same pipeline: IV screens features, "
    "WoE transforms them, and logistic regression combines them into a score.")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION F — How it influenced OUR project
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "F.", "How It Influenced Our Project")
add_rule(doc)

add_body(doc,
    "Bivariate analysis was the second analytical step in our pipeline, applied after univariate "
    "cleaning and transformation. Its outputs directly determined which features entered the "
    "logistic regression and how they were encoded:")

influences = [
    ("IV-based feature short-listing. ",
     "We computed IV for all 23 candidate features against the binary default flag. Features with "
     "IV < 0.02 were excluded from further analysis: 6 features were dropped at this stage "
     "(including loan currency, product flag, and two date-derived variables with no discriminatory "
     "power). This reduced the modelling feature set from 23 to 17, lowering the risk of overfitting "
     "and reducing scorecard complexity for regulatory review."),
    ("WoE transformation for all retained features. ",
     "The 17 surviving features were WoE-transformed using monotonically constrained bins. "
     "The monotonicity constraint (WoE must be strictly increasing or decreasing across bins) "
     "was enforced algorithmically and validated visually — a non-monotonic WoE pattern indicates "
     "an unstable or overfitted binning that would produce unreliable scorecard weights."),
    ("Multicollinearity detection and feature pruning. ",
     "A Spearman correlation matrix of all 17 WoE-transformed features revealed two strongly "
     "correlated pairs: (a) log(income) and DTI ratio (ρ = −0.61, borderline) and (b) "
     "number of open accounts and total credit exposure (ρ = 0.74, exceeding the threshold). "
     "For pair (b), total credit exposure was retained (higher IV = 0.29 vs. 0.17) and open "
     "accounts was dropped, reducing multicollinearity before logistic regression fitting."),
    ("Chi-square validation for categorical features. ",
     "For the five categorical features (employment type, province, loan purpose, marital status, "
     "housing type), chi-square tests confirmed statistical association with default (all p < 0.001 "
     "after Bonferroni correction). Cramér's V guided the grouping of rare categories — any "
     "category with fewer than 200 observations was merged with its nearest WoE neighbour."),
    ("Suspicious IV investigation. ",
     "One feature (days since last delinquency) returned IV = 0.61, flagged as suspicious "
     "under the IV > 0.50 rule. Investigation revealed that this feature used event date data "
     "that would not be available at loan origination — a form of look-ahead data leakage. "
     "The feature was removed from the pipeline entirely. Bivariate IV analysis was the mechanism "
     "that surfaced this leakage before it could contaminate the model."),
    ("WoE monotonicity as a model constraint signal. ",
     "The WoE analysis for DTI ratio showed a strong, clean monotonic pattern (WoE increases "
     "steadily with DTI). This gave us confidence that a linear logistic regression (which assumes "
     "monotonicity after WoE transformation) is appropriate for this variable — no additional "
     "non-linear terms or interaction features were needed."),
]
for bold, rest in influences:
    add_bullet(doc, rest, bold_prefix=bold)

add_rule(doc)
footer = doc.add_paragraph(
    "FNB DataQuest 2026  ·  Credit Risk Modelling Team  ·  Confidential & Internal Use Only")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), "Bivariate_Analysis_Research_Notes.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
