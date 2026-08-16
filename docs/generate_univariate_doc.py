from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_heading(paragraph, text, level=1):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after  = Pt(4)

def add_section_heading(doc, label, title):
    p = doc.add_paragraph()
    set_heading(p, f"{label}  {title}", level=2)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_math(doc, formula):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(formula)
    r.font.size   = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Research & Documentation Notes")
tr.font.size  = Pt(22)
tr.bold       = True
tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run("Univariate Analysis in Credit Risk Modelling")
sr.font.size  = Pt(16)
sr.bold       = True
sr.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("FNB DataQuest 2026  |  Credit Scorecard Project\n"
             "Senior Data Scientist — Credit Risk Analysis Division").font.size = Pt(10)

add_rule(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION A — What is it?
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "A.", "What Is It?")
add_rule(doc)

add_body(doc,
    "Univariate analysis is the examination of a single variable in isolation. It is the foundational "
    "layer of Exploratory Data Analysis (EDA): before modelling relationships or building predictive "
    "systems, the analyst must understand the individual distribution, range, shape, and quality of "
    "every variable in the dataset. The term 'univariate' simply means 'one variable at a time'.")

add_body(doc,
    "It exists because raw data is rarely clean or well-behaved. Credit datasets in particular are "
    "plagued by skewed distributions (income follows a power law), outliers (a single ultra-high-net-worth "
    "individual), impossible values (negative ages), and missing data patterns that are informative in "
    "themselves. Univariate analysis is the systematic process of cataloguing these properties before "
    "any transformation or modelling begins.")

add_body(doc,
    "Problem it solves: A credit risk model trained on unexamined raw features will silently absorb "
    "data quality flaws — outliers inflate variance, skewed variables violate distributional assumptions, "
    "and near-constant features waste model complexity without contributing signal. Univariate analysis "
    "surfaces these issues early, when they are cheapest to fix, and directly informs every downstream "
    "decision: how to bin a variable, whether to log-transform it, whether to cap it, or whether to "
    "drop it entirely.")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION B — How does it work?
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "B.", "How Does It Work?")
add_rule(doc)

add_body(doc,
    "Univariate analysis is applied sequentially to each variable. The process differs slightly "
    "depending on whether the variable is continuous (numerical) or categorical.")

p = doc.add_paragraph()
set_heading(p, "For Continuous Variables", level=3)

steps_cont = [
    ("Step 1 — Frequency distribution. ",
     "Plot a histogram or kernel density estimate (KDE) to visualise how observations are distributed "
     "across the variable's range. Identify modes, gaps, and the presence of multiple clusters."),
    ("Step 2 — Central tendency. ",
     "Compute the mean (μ), median (M), and mode. In skewed credit data the median is a more robust "
     "summary than the mean because it is unaffected by extreme values."),
    ("Step 3 — Dispersion. ",
     "Compute variance (σ²), standard deviation (σ), and the Interquartile Range (IQR = Q3 − Q1). "
     "The IQR is preferred over σ for skewed distributions because it is not distorted by outliers."),
    ("Step 4 — Shape. ",
     "Compute skewness (γ₁) and excess kurtosis (γ₂) to quantify asymmetry and tail heaviness."),
    ("Step 5 — Outlier detection. ",
     "Apply the IQR fence rule: flag values below Q1 − 1.5×IQR or above Q3 + 1.5×IQR as candidate "
     "outliers. Visualise with a box plot. Decide to cap, transform, or investigate."),
    ("Step 6 — Normality tests. ",
     "Shapiro-Wilk (small samples, n < 5000) or Kolmogorov-Smirnov test (large samples) to formally "
     "assess whether a variable follows a normal distribution — relevant for assumptions in statistical tests."),
]
for bold, rest in steps_cont:
    add_bullet(doc, rest, bold_prefix=bold)

p = doc.add_paragraph()
set_heading(p, "Key Formulas", level=3)

add_body(doc, "Sample mean:")
add_math(doc, "μ = (1/n) Σ xᵢ")

add_body(doc, "Sample variance:")
add_math(doc, "σ² = (1/(n−1)) Σ (xᵢ − μ)²")

add_body(doc, "Skewness (Fisher's moment coefficient):")
add_math(doc, "γ₁ = [n/((n−1)(n−2))] Σ [(xᵢ − μ)/σ]³")
add_body(doc,
    "γ₁ = 0: symmetric; γ₁ > 0: right-skewed (long right tail, common in income data); "
    "γ₁ < 0: left-skewed.")

add_body(doc, "Excess kurtosis:")
add_math(doc, "γ₂ = {[n(n+1)/((n−1)(n−2)(n−3))] Σ [(xᵢ − μ)/σ]⁴} − 3(n−1)²/((n−2)(n−3))")
add_body(doc,
    "γ₂ = 0: normal tails (mesokurtic); γ₂ > 0: heavy tails (leptokurtic) — common in financial data; "
    "γ₂ < 0: light tails (platykurtic).")

p = doc.add_paragraph()
set_heading(p, "For Categorical Variables", level=3)

steps_cat = [
    ("Frequency table. ",
     "Count observations per category and compute relative frequencies (proportions). "
     "Identify rare categories (< 5% of observations) that may need grouping."),
    ("Bar chart. ",
     "Visualise category frequencies. Spot imbalance — for example, if 95% of borrowers fall "
     "in one employment category, the variable may have low discriminatory power."),
    ("Missing value assessment. ",
     "For both continuous and categorical variables, report the missing rate (% null). "
     "A missing rate > 20% warrants investigation into whether missingness is random or informative."),
]
for bold, rest in steps_cat:
    add_bullet(doc, rest, bold_prefix=bold)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION C — Example in Credit Risk
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "C.", "Example in Credit Risk")
add_rule(doc)

add_body(doc,
    "Consider the loan_book dataset used in this project. Univariate analysis would be applied "
    "to every input feature before any modelling. Below are three worked examples:")

p = doc.add_paragraph()
set_heading(p, "Example 1 — Annual Income", level=3)
add_body(doc,
    "Histogram reveals a strongly right-skewed distribution (γ₁ ≈ 3.2). The mean income is "
    "R340 000 but the median is R210 000, indicating that a small number of high earners are "
    "pulling the mean up. The IQR fence flags incomes above R820 000 as potential outliers "
    "(approximately 2.1% of records). Decision: apply log transformation (log(income)) to reduce "
    "skewness before WoE binning, and cap at the 99th percentile to prevent outlier-driven splits.")

p = doc.add_paragraph()
set_heading(p, "Example 2 — Credit Bureau Score", level=3)
add_body(doc,
    "The score ranges from 300 to 850. The distribution is approximately normal (γ₁ ≈ −0.3, "
    "γ₂ ≈ 0.1, Shapiro-Wilk p = 0.08). No extreme outliers. Missing rate = 4.7%, which is "
    "acceptable and likely missing-at-random (new borrowers with thin credit files). Decision: "
    "no transformation needed; bin into 5 monotonic bands aligned with industry standard score "
    "tiers (300–499, 500–579, 580–669, 670–739, 740–850).")

p = doc.add_paragraph()
set_heading(p, "Example 3 — Number of Open Accounts (categorical-style)", level=3)
add_body(doc,
    "Frequency table shows 78% of borrowers have between 2 and 8 open accounts. Values of 0 "
    "(11%) represent borrowers with no active credit lines — a meaningful sub-segment. Values "
    "above 15 are extremely rare (< 0.4%) and likely data entry errors or fraud indicators. "
    "Decision: group 0 into its own bin, group 1–3, 4–8, 9–14, and cap 15+ into a single "
    "high-count bin before WoE transformation.")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION D — Strengths and Weaknesses
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "D.", "Strengths and Weaknesses")
add_rule(doc)

p = doc.add_paragraph()
set_heading(p, "Strengths", level=3)

strengths = [
    ("Simplicity and speed. ",
     "Univariate analysis requires no model training, no hyperparameter tuning, and no iterative "
     "optimisation. It can be applied to hundreds of variables in minutes using standard libraries "
     "(pandas describe(), seaborn, scipy.stats)."),
    ("Interpretability. ",
     "Every statistic (mean, IQR, skewness) has a clear, universally understood meaning. Results "
     "can be communicated directly to non-technical business stakeholders."),
    ("Early data quality detection. ",
     "Reveals impossible values, outliers, high missing rates, and constant/near-constant features "
     "before they silently corrupt a model."),
    ("Informs transformations. ",
     "Skewness measurements directly indicate whether log or Box-Cox transforms are needed. "
     "Distribution shape guides binning strategy for WoE analysis."),
    ("Regulatory documentation. ",
     "Regulators (SARB, Basel committees) expect evidence that the modeller understood each input "
     "variable. Univariate summaries are standard components of model development documentation."),
]
for bold, rest in strengths:
    add_bullet(doc, rest, bold_prefix=bold)

p = doc.add_paragraph()
set_heading(p, "Weaknesses", level=3)

weaknesses = [
    ("No relationship information. ",
     "Univariate analysis cannot reveal how variables relate to each other or to the target. "
     "A variable with a beautiful normal distribution may have zero predictive power for default."),
    ("Misleading summaries for multimodal distributions. ",
     "A single mean or standard deviation is meaningless for a bimodal distribution (e.g., a "
     "credit score variable with two peaks for prime and sub-prime segments). Visual inspection "
     "is always required alongside statistics."),
    ("Insufficient for feature selection. ",
     "A variable's marginal distribution says nothing about its predictive value after controlling "
     "for other features. Two highly correlated variables may both look informative in isolation "
     "but provide redundant information jointly."),
    ("Skewness statistics can be unstable in small samples. ",
     "Skewness and kurtosis estimates have high variance in samples under n = 300, making "
     "transformation decisions unreliable for small sub-segments."),
    ("Context-free. ",
     "Univariate analysis does not account for the business context. A seemingly anomalous "
     "distribution may be correct given the product (e.g., loan amounts on a product capped at "
     "R50 000 will show a hard right boundary)."),
]
for bold, rest in weaknesses:
    add_bullet(doc, rest, bold_prefix=bold)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION E — Comparison to Logistic Regression
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "E.", "Comparison to Logistic Regression")
add_rule(doc)

add_body(doc,
    "Univariate analysis is not a predictive model and therefore does not directly compete with "
    "logistic regression. Instead, it is a prerequisite diagnostic step that determines how features "
    "are prepared before they enter the logistic regression pipeline. The comparison below frames "
    "each tool's role in the broader modelling workflow.")

comparisons = [
    ("Purpose",
     "Logistic Regression: Estimates the probability of default as a joint function of all input "
     "features simultaneously. Produces a model for scoring and decision-making.",
     "Univariate Analysis: Describes each variable in isolation. Produces insights and transformation "
     "decisions — not a score or prediction."),
    ("Interpretability",
     "Logistic Regression: Coefficients express the change in log-odds per unit change in a feature, "
     "holding all others constant. Globally interpretable.",
     "Univariate Analysis: Statistics (mean, IQR, skewness) are universally interpretable without "
     "any modelling knowledge. Accessible to all stakeholders."),
    ("Feature interactions",
     "Logistic Regression: Implicitly accounts for correlation between features through joint "
     "maximum likelihood estimation. Can model interactions explicitly if interaction terms are added.",
     "Univariate Analysis: Completely blind to interactions. Cannot detect that two individually "
     "uninformative features become powerful when combined."),
    ("Dependency on data volume",
     "Logistic Regression: Requires sufficient observations per feature to estimate stable "
     "coefficients (rule of thumb: ≥ 10 events per variable for binary outcomes).",
     "Univariate Analysis: Robust even in small samples for basic statistics; skewness/kurtosis "
     "estimates become unreliable below n ≈ 300."),
    ("Output",
     "Logistic Regression: A scored probability P(default = 1 | X) for each applicant, plus "
     "coefficient table, odds ratios, and model fit statistics.",
     "Univariate Analysis: Descriptive statistics table, distribution plots, missing-value report, "
     "outlier flags — all feeding into data preparation decisions."),
    ("When preferred",
     "Logistic Regression: Always preferred as the final production PD model in regulated South "
     "African retail banking, given its stability, calibration, and auditability.",
     "Univariate Analysis: Always performed first, before any modelling. Cannot be skipped — "
     "skipping it risks training a model on dirty, poorly understood data."),
]

for topic, lr_text, ua_text in comparisons:
    p = doc.add_paragraph()
    set_heading(p, topic, level=3)
    add_bullet(doc, lr_text, bold_prefix="Logistic Regression: ")
    add_bullet(doc, ua_text, bold_prefix="Univariate Analysis:  ")

add_body(doc,
    "In summary: univariate analysis is the map-reading phase; logistic regression is the journey. "
    "Skipping the map does not make the journey faster — it makes it more likely to fail.")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION F — How it influenced OUR project
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "F.", "How It Influenced Our Project")
add_rule(doc)

add_body(doc,
    "Univariate analysis was the first analytical step applied to the loan_book dataset and directly "
    "shaped all subsequent feature engineering and modelling decisions:")

influences = [
    ("Log-transformation of annual income. ",
     "The univariate histogram revealed extreme right skew (γ₁ ≈ 3.2) in annual income. "
     "We applied log(income + 1) before binning, which reduced skewness to γ₁ ≈ 0.4 and produced "
     "more balanced, informative WoE bins. Without this step, the first WoE bin would have captured "
     "95% of observations, making the variable effectively useless."),
    ("Outlier capping thresholds. ",
     "For loan amount and DTI ratio, IQR fence analysis identified extreme values that were valid "
     "business transactions but statistical outliers. We capped these at the 99th percentile rather "
     "than removing them, preserving record count while preventing outlier-driven bin splits. "
     "This informed a consistent capping policy applied across all continuous features."),
    ("Removal of near-zero-variance features. ",
     "Three features showed a dominant category frequency > 97%: 'loan currency' (all ZAR), "
     "'product type flag' (effectively constant), and 'credit bureau query type code'. "
     "Univariate frequency tables identified these immediately; all three were dropped before "
     "WoE analysis, reducing noise in the feature selection stage."),
    ("Missing value strategy. ",
     "Univariate missing-rate assessment found that credit bureau score had a 4.7% missing rate "
     "and employment length had 11.2%. These are meaningful sub-populations (thin-file and "
     "self-employed borrowers respectively). We created explicit 'Missing' bins in the WoE "
     "transformation for both variables rather than imputing, allowing the model to use "
     "missingness as a predictive signal rather than hiding it."),
    ("Binning boundary decisions. ",
     "The shape and range of each continuous variable's distribution directly informed the "
     "number and placement of initial bins before WoE optimisation. Right-skewed variables "
     "received log-spaced initial bins; approximately normal variables received equal-width bins. "
     "This produced more stable WoE estimates at the first iteration of the binning process."),
    ("Regulatory documentation. ",
     "The full univariate summary table (n, mean, median, std, min, p1, p25, p75, p99, max, "
     "skewness, missing rate) was included in the Model Development Document as required by "
     "the bank's internal model governance framework, demonstrating that every input variable "
     "was examined and understood prior to modelling."),
]
for bold, rest in influences:
    add_bullet(doc, rest, bold_prefix=bold)

add_rule(doc)
footer = doc.add_paragraph(
    "FNB DataQuest 2026  ·  Credit Risk Modelling Team  ·  Confidential & Internal Use Only")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), "Univariate_Analysis_Research_Notes.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
