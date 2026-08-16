from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_heading(paragraph, text, level=1):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after  = Pt(4)

def add_section_heading(doc, label, title):
    p = doc.add_paragraph()
    set_heading(p, f"{label}  {title}", level=2)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_math(doc, formula):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(formula)
    r.font.size   = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Research & Documentation Notes")
tr.font.size  = Pt(22)
tr.bold       = True
tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run("Classification Trees in Credit Risk Modelling")
sr.font.size  = Pt(16)
sr.bold       = True
sr.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("FNB DataQuest 2026  |  Credit Scorecard Project\n"
             "Senior Data Scientist — Credit Risk Analysis Division").font.size = Pt(10)

add_rule(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION A — What is it?
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "A.", "What Is It?")
add_rule(doc)

add_body(doc,
    "A Classification Tree (also known as a Decision Tree Classifier) is a supervised machine learning "
    "algorithm that partitions the feature space into non-overlapping rectangular regions using a series "
    "of recursive binary splits. At each internal node the tree asks a simple yes/no question about a "
    "single feature (e.g., 'Is annual income ≥ R250 000?'), routes observations left or right depending "
    "on the answer, and eventually assigns every observation in a terminal leaf to a majority class label.")

add_body(doc,
    "The method exists because real-world classification boundaries are rarely linear. Logistic regression "
    "draws a single hyperplane through feature space; a tree instead builds a staircase-like boundary that "
    "can approximate complex, non-linear relationships without any algebraic transformation of the inputs. "
    "Crucially, the resulting structure mirrors a human analyst's rule-based thinking — a chain of "
    "'if–then–else' conditions — making it highly interpretable to business stakeholders and regulators.")

add_body(doc,
    "Problem it solves: In credit risk, the decision to approve or decline a loan application depends on "
    "the interplay of many borrower characteristics. Some relationships are non-linear (e.g., default risk "
    "spikes sharply only above a debt-to-income threshold), and some features interact with each other "
    "(e.g., low income is only dangerous when combined with high existing debt). Classification trees "
    "discover these thresholds and interactions automatically from data, without requiring the analyst to "
    "specify them in advance.")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION B — How does it work?
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "B.", "How Does It Work?")
add_rule(doc)

add_body(doc,
    "Classification trees are built using a top-down, greedy recursive partitioning algorithm (most "
    "commonly CART — Classification And Regression Trees, Breiman et al. 1984). The steps are:")

steps = [
    ("Step 1 — Start at the root node. ", "The entire training dataset sits at the root."),
    ("Step 2 — Search for the best split. ",
     "For every candidate feature j and every candidate threshold t, compute the impurity reduction "
     "produced by splitting the node into two child sets:  left = {x : xⱼ ≤ t}  and  right = {x : xⱼ > t}."),
    ("Step 3 — Select the split that maximises impurity reduction. ",
     "The winning (j*, t*) is the pair that produces the largest decrease in impurity."),
    ("Step 4 — Recurse. ",
     "Apply steps 2–3 independently to each child node. Continue until a stopping criterion is met "
     "(maximum depth, minimum samples per leaf, or negligible impurity gain)."),
    ("Step 5 — Assign class labels at leaves. ",
     "Each terminal leaf is assigned the majority class of training observations that fall into it. "
     "The predicted probability is the empirical class proportion at that leaf."),
    ("Step 6 — Prune (optional but recommended). ",
     "Post-pruning (cost-complexity pruning) removes branches that do not improve performance on a "
     "validation set, reducing overfitting."),
]
for bold, rest in steps:
    add_bullet(doc, rest, bold_prefix=bold)

# Impurity measures
p = doc.add_paragraph()
set_heading(p, "Key Impurity Measures", level=3)

add_body(doc,
    "The most common impurity measures used at each node are Gini Impurity and Entropy (Information Gain).")

add_body(doc, "Gini Impurity (used by scikit-learn CART by default):")
add_math(doc, "Gini(t) = 1 − Σ pₖ²   (summed over classes k at node t)")
add_body(doc,
    "where pₖ is the proportion of class k observations at node t. A pure node (all one class) has "
    "Gini = 0; a perfectly mixed binary node has Gini = 0.5.")

add_body(doc, "Entropy (Information Gain):")
add_math(doc, "H(t) = − Σ pₖ log₂(pₖ)")
add_body(doc,
    "Information Gain at a split = H(parent) − [weighted average of H(left) + H(weighted average H(right))].")

add_body(doc, "Impurity Reduction (the criterion maximised at each node):")
add_math(doc,
    "ΔImpurity = Impurity(parent) − [N_L/N · Impurity(left)] − [N_R/N · Impurity(right)]")

add_body(doc,
    "Key assumptions: (1) splits are axis-aligned (one feature at a time); "
    "(2) the algorithm is greedy — it optimises each split locally, not globally; "
    "(3) all input types (continuous, ordinal, nominal) can be handled with appropriate encoding; "
    "(4) missing values can be handled via surrogate splits or imputation.")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION C — Example in Credit Risk
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "C.", "Example in Credit Risk")
add_rule(doc)

add_body(doc,
    "Consider a retail bank processing personal loan applications. The target variable is binary: "
    "Default (1) vs. No Default (0). Features include annual income, credit bureau score, "
    "debt-to-income ratio (DTI), number of delinquencies in the past 24 months, employment length, "
    "and loan amount requested.")

add_body(doc,
    "A classification tree trained on this data might produce a path such as:")

tree_rules = [
    "Root: Is credit bureau score < 580?  →  YES → HIGH-RISK branch",
    "    |→ Is DTI ratio > 0.45?  →  YES → Leaf: Predict DEFAULT (p_default = 0.81)",
    "    |→ Is DTI ratio ≤ 0.45?  →  Is number of delinquencies > 2?",
    "         |→ YES → Leaf: Predict DEFAULT (p_default = 0.67)",
    "         |→ NO  → Leaf: Predict NO DEFAULT (p_default = 0.31)",
    "Root: Is credit bureau score ≥ 580?  →  LOW-RISK branch",
    "    |→ Is annual income < R180 000?  →  YES → Is loan amount > R60 000?",
    "         |→ YES → Leaf: Predict DEFAULT (p_default = 0.52)",
    "         |→ NO  → Leaf: Predict NO DEFAULT (p_default = 0.18)",
    "    |→ Is annual income ≥ R180 000?  →  Leaf: Predict NO DEFAULT (p_default = 0.07)",
]
for rule in tree_rules:
    p = doc.add_paragraph(rule, style="List Bullet")
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.name = "Courier New"

add_body(doc,
    "This tree structure is immediately interpretable to a credit analyst or regulator: each path from "
    "root to leaf is a lending policy rule. A borrower with credit score 540 and DTI 0.50 is routed to "
    "the first leaf (p_default = 0.81) and would be declined. A borrower with score 620 and income "
    "R200 000 is routed to the last leaf (p_default = 0.07) and would be approved.")

add_body(doc,
    "The bank can also extract feature importance rankings directly from the tree — features that appear "
    "higher in the tree and produce larger impurity reductions are more predictive of default, helping "
    "credit policy teams prioritise which variables to monitor and report to regulators.")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION D — Strengths and Weaknesses
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "D.", "Strengths and Weaknesses")
add_rule(doc)

p = doc.add_paragraph()
set_heading(p, "Strengths", level=3)

strengths = [
    ("Interpretability. ",
     "The tree structure maps directly to human-readable if-then rules. A single-page diagram can "
     "communicate the full model to a credit committee without statistical expertise."),
    ("Handles non-linearity and interactions automatically. ",
     "No manual feature engineering is needed to capture threshold effects or interaction terms."),
    ("No scaling required. ",
     "Because splits use rank orderings, features do not need to be standardised or normalised."),
    ("Handles mixed data types. ",
     "Categorical and continuous variables can both be used as splitting criteria without encoding tricks."),
    ("Transparent probability estimates. ",
     "Leaf-level default rates are directly observable empirical proportions, not black-box outputs."),
    ("Fast to score. ",
     "Prediction at inference time is a sequence of binary comparisons — O(depth) — making real-time "
     "decisioning highly efficient."),
]
for bold, rest in strengths:
    add_bullet(doc, rest, bold_prefix=bold)

p = doc.add_paragraph()
set_heading(p, "Weaknesses", level=3)

weaknesses = [
    ("High variance / overfitting. ",
     "An unconstrained tree memorises training data. A small perturbation in the training set can produce "
     "a completely different tree structure. Pruning and depth constraints help, but the problem is fundamental."),
    ("Greedy splitting is suboptimal. ",
     "Local optimisation at each node does not guarantee a globally optimal tree. A different first split "
     "might produce a better tree overall."),
    ("Instability. ",
     "Trees are sensitive to outliers and small data changes. This undermines reproducibility — a concern "
     "in regulatory model validation."),
    ("Poor calibration of probability estimates. ",
     "Leaf probabilities are based on potentially small sample counts, making them noisy. Calibration "
     "post-processing (Platt scaling, isotonic regression) is often required."),
    ("Biased toward high-cardinality features. ",
     "When using Gini or Entropy without correction, features with many unique values have more candidate "
     "split points and appear more informative by chance."),
    ("Regulatory limitations. ",
     "A single unstable tree is harder to defend in Basel/IFRS 9 model governance than a logistic "
     "regression with stable, auditable coefficients."),
    ("Axis-aligned boundaries. ",
     "The staircase decision boundary is inefficient for truly diagonal or elliptical class boundaries."),
]
for bold, rest in weaknesses:
    add_bullet(doc, rest, bold_prefix=bold)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION E — Comparison to Logistic Regression
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "E.", "Comparison to Logistic Regression")
add_rule(doc)

comparison = [
    ("Interpretability",
     "Logistic Regression",
     "Linear coefficients as log-odds; each coefficient has a single, globally consistent meaning. "
     "Preferred by regulators (Basel III, IFRS 9, SR 11-7).",
     "Classification Tree",
     "If-then rule chains; interpretable locally (per path) but no single global equation. "
     "Harder to defend coefficient-style attribution."),
    ("Non-linearity",
     "Logistic Regression",
     "Requires manual engineering (splines, bins, interaction terms) to capture non-linearity.",
     "Classification Tree",
     "Captures non-linearity and interactions automatically through recursive splitting."),
    ("Performance",
     "Logistic Regression",
     "Strong baseline; competitive on linearly separable problems and well-engineered features.",
     "Classification Tree",
     "Can outperform logistic regression on datasets with complex interactions, but a single tree "
     "typically underperforms regularised logistic regression on tabular credit data."),
    ("Stability",
     "Logistic Regression",
     "Coefficients are stable under small data perturbations (especially with L2 regularisation).",
     "Classification Tree",
     "High variance; structure can change dramatically with small training set changes."),
    ("Calibration",
     "Logistic Regression",
     "Probabilities are well-calibrated by construction (sigmoid output trained with log-loss).",
     "Classification Tree",
     "Leaf probabilities can be noisy; requires post-processing for reliable PD estimates."),
    ("Feature handling",
     "Logistic Regression",
     "Requires encoding, scaling, and often WoE transformation for optimal performance.",
     "Classification Tree",
     "Handles raw mixed types; no scaling required."),
    ("Regulatory acceptance",
     "Logistic Regression",
     "Industry standard for PD modelling under Basel frameworks. Auditable and explainable.",
     "Classification Tree",
     "Accepted in some contexts but less common as a standalone PD model in regulated environments."),
]

for topic, lr_label, lr_text, ct_label, ct_text in comparison:
    p = doc.add_paragraph()
    set_heading(p, topic, level=3)
    add_bullet(doc, lr_text, bold_prefix="Logistic Regression: ")
    add_bullet(doc, ct_text, bold_prefix="Classification Tree:  ")

add_body(doc,
    "When to prefer Logistic Regression: When regulatory explainability is paramount; when probability "
    "calibration is critical (e.g., IFRS 9 PD term structures); when the relationship between features "
    "and log-odds is approximately linear after WoE transformation; when model stability and reproducibility "
    "are required for governance sign-off.")

add_body(doc,
    "When Classification Trees are superior: When exploring non-linear interactions for feature engineering "
    "insight; when building an ensemble base learner (Random Forest, Gradient Boosting); when generating "
    "segmentation rules or policy scorecards for operational teams; when rapid prototyping is needed without "
    "feature transformation.")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION F — How it influenced OUR project
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "F.", "How It Influenced Our Project")
add_rule(doc)

add_body(doc,
    "While our primary model is a logistic regression scorecard, the study of classification trees "
    "directly informed several design decisions across our pipeline:")

influences = [
    ("Discovering non-linear thresholds for binning. ",
     "We used a shallow classification tree (max_depth=4) as an exploratory tool to identify natural "
     "breakpoints in continuous features such as DTI ratio, credit bureau score, and loan-to-value ratio. "
     "The split thresholds identified by the tree (e.g., DTI > 0.43, bureau score < 590) were used to "
     "inform our manual binning strategy before WoE transformation, replacing arbitrary quantile-based "
     "cuts with data-driven, business-meaningful boundaries."),
    ("Identifying interaction terms. ",
     "By examining the tree's branching logic (e.g., income is only predictive conditional on credit "
     "score being below a threshold), we identified two-way interaction features to include in our "
     "logistic regression — specifically an income × delinquency interaction term that significantly "
     "improved the model's AUC."),
    ("Feature importance screening. ",
     "We used Gini-based feature importance from a classification tree as a fast pre-screening step to "
     "identify candidate variables before running WoE IV analysis. Features with near-zero tree importance "
     "were deprioritised for deeper analysis, reducing computation time."),
    ("Segment-specific sub-models. ",
     "The tree's natural partitioning suggested two borrower segments (prime vs. sub-prime, defined by "
     "bureau score ≥ 580) with meaningfully different default dynamics. We explored segment-specific "
     "logistic regression models for these sub-populations, which showed improved Gini coefficients "
     "versus a single pooled model."),
    ("Avoiding over-reliance on complex models. ",
     "Studying the high variance and poor calibration of single classification trees reinforced our "
     "decision to use logistic regression as the final production model. The interpretability and "
     "stability requirements of a regulated South African retail banking environment make logistic "
     "regression the appropriate and defensible choice for our scorecard."),
    ("Calibration awareness. ",
     "Understanding that tree leaf probabilities are poorly calibrated motivated us to apply Platt "
     "scaling in our ensemble experiments and to verify probability calibration plots for our final "
     "logistic regression model, ensuring reliable PD estimates for IFRS 9 provisioning."),
]
for bold, rest in influences:
    add_bullet(doc, rest, bold_prefix=bold)

add_rule(doc)
footer = doc.add_paragraph(
    "FNB DataQuest 2026  ·  Credit Risk Modelling Team  ·  Confidential & Internal Use Only")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), "Classification_Trees_Research_Notes.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
